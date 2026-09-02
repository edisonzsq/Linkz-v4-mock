# -*- coding: utf-8 -*-
"""
Shared builder for the "open questions" answer sheets sent to the spec author.

One round of questions = one small script that imports `Sheet` from here and
calls `question()` / `para()` / `option()` / `answer_box()`. See
`build-questions-docx.py` (round 1) and `build-questions-2-docx.py` (round 2).

Two things this exists to get right, both of which produce a file that
python-docx will happily reopen but Word and LibreOffice refuse to load:

  - **OOXML child order is strict.** `w:shd`, `w:tcBorders`, `w:tcMar` and
    friends must appear in schema sequence inside `tcPr` / `pPr` / `tblPr`.
    Use `insert_ordered()`, never `parent.append()`.
  - **Column widths need `set_widths()`.** Setting `cell.width` alone does
    nothing; the table also needs a fixed `w:tblLayout` and a matching
    `w:tblGrid`.

Check any change by converting and looking at it — `libreoffice-core` alone
cannot open .docx, you need `libreoffice-writer`:

    soffice --headless --convert-to pdf --outdir /tmp/out "<file>.docx"
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
RED = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x60, 0x60, 0x60)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)


# ---------- OOXML schema ordering (LibreOffice/Word reject out-of-order children) ----------
_ORDERS = {
    'tcPr': ['w:cnfStyle', 'w:tcW', 'w:gridSpan', 'w:hMerge', 'w:vMerge', 'w:tcBorders',
             'w:shd', 'w:noWrap', 'w:tcMar', 'w:textDirection', 'w:tcFitText', 'w:vAlign',
             'w:hideMark'],
    'pPr': ['w:pStyle', 'w:keepNext', 'w:keepLines', 'w:pageBreakBefore', 'w:framePr',
            'w:widowControl', 'w:numPr', 'w:suppressLineNumbers', 'w:pBdr', 'w:shd', 'w:tabs',
            'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap', 'w:overflowPunct',
            'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd',
            'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents',
            'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
            'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr',
            'w:sectPr', 'w:pPrChange'],
    'tblPr': ['w:tblStyle', 'w:tblpPr', 'w:tblOverlap', 'w:bidiVisual', 'w:tblStyleRowBandSize',
              'w:tblStyleColBandSize', 'w:tblW', 'w:jc', 'w:tblCellSpacing', 'w:tblInd',
              'w:tblBorders', 'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook',
              'w:tblCaption', 'w:tblDescription', 'w:tblPrChange'],
    'rPr': ['w:rStyle', 'w:rFonts', 'w:b', 'w:bCs', 'w:i', 'w:iCs', 'w:caps', 'w:smallCaps',
            'w:strike', 'w:dstrike', 'w:outline', 'w:shadow', 'w:emboss', 'w:imprint',
            'w:noProof', 'w:snapToGrid', 'w:vanish', 'w:webHidden', 'w:color', 'w:spacing',
            'w:w', 'w:kern', 'w:position', 'w:sz', 'w:szCs', 'w:highlight', 'w:u', 'w:effect',
            'w:bdr', 'w:shd', 'w:fitText', 'w:vertAlign', 'w:rtl', 'w:cs', 'w:em', 'w:lang',
            'w:eastAsianLayout', 'w:specVanish', 'w:oMath'],
}


def _local(tag):
    return tag.split('}')[-1]


def insert_ordered(parent, child):
    """Insert child into parent at its schema-mandated position."""
    order = _ORDERS.get(_local(parent.tag))
    name = 'w:' + _local(child.tag)
    if order is None or name not in order:
        parent.append(child)
        return child
    idx = order.index(name)
    for existing in parent:
        en = 'w:' + _local(existing.tag)
        if en in order and order.index(en) > idx:
            existing.addprevious(child)
            return child
    parent.append(child)
    return child


# ---------- low-level helpers ----------
def shade(el, hexfill):
    for old in el.findall(qn('w:shd')):
        el.remove(old)
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hexfill)
    insert_ordered(el, sh)


def cell_shade(cell, hexfill):
    shade(cell._tc.get_or_add_tcPr(), hexfill)


def cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        e = OxmlElement('w:' + tag)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    insert_ordered(tcPr, mar)


def borders(el, color='BFBFBF', sz=6, edges=('top', 'left', 'bottom', 'right'), val='single'):
    """el is a tcPr or pPr-borders parent; returns the created element name used."""
    tag = 'w:tcBorders' if _local(el.tag) == 'tcPr' else 'w:pBdr'
    bd = el.find(qn(tag))
    if bd is None:
        bd = OxmlElement(tag)
        insert_ordered(el, bd)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge not in edges:
            continue
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), val)
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '4' if tag == 'w:pBdr' else '0')
        e.set(qn('w:color'), color)
        bd.append(e)


def cell_borders(cell, color='BFBFBF', sz=6, edges=('top', 'left', 'bottom', 'right')):
    borders(cell._tc.get_or_add_tcPr(), color, sz, edges)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    bd = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        bd.append(e)
    insert_ordered(tblPr, bd)


def run(p, text, bold=False, italic=False, color=None, size=None, font=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    if font is not None:
        r.font.name = font
        r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return r


def rich(p, parts):
    """parts: list of (text, bold, italic) or plain str."""
    for part in parts:
        if isinstance(part, str):
            run(p, part)
        else:
            text, bold, italic = (list(part) + [False, False])[:3]
            run(p, text, bold=bold, italic=italic)
    return p


def _para(doc, parts, space_after=8, indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    rich(p, parts if isinstance(parts, list) else [parts])
    return p


def _quote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.28)
    pf.right_indent = Inches(0.2)
    pf.space_before = Pt(6)
    pf.space_after = Pt(10)
    borders(p._p.get_or_add_pPr(), color='9DB2CE', sz=18, edges=('left',))
    shade(p._p.get_or_add_pPr(), 'F2F5FA')
    run(p, text, italic=True, color=GREY)
    return p


def _bullet(doc, parts, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + 0.28 * level)
    p.paragraph_format.space_after = Pt(4)
    rich(p, parts if isinstance(parts, list) else [parts])
    return p


def _numbered(doc, parts):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(6)
    rich(p, parts if isinstance(parts, list) else [parts])
    return p


def _option(doc, label, rest='', level=0):
    """A tickable option line: ☐  **label** rest"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.32 + 0.3 * level)
    pf.first_line_indent = Inches(-0.32)
    pf.space_after = Pt(5)
    run(p, '☐', size=13, font='Segoe UI Symbol')
    run(p, '   ')
    run(p, label, bold=True)
    if rest:
        run(p, rest)
    return p


def _rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    borders(p._p.get_or_add_pPr(), color='D0D0D0', sz=6, edges=('bottom',))
    return p


def _answer_box(doc, hint='Click here and type your answer.', lines=4):
    # Pull the last line of the question onto the same page as its box, so the
    # box is never the only thing at the top of a page.
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.keep_with_next = True

    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_widths(t, [6.45])

    head = t.rows[0].cells[0]
    head.text = ''
    cell_shade(head, 'FFE9A8')
    cell_borders(head, color='D9A400', sz=8)
    cell_margins(head, top=60, bottom=60)
    hp = head.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    # Keep the header strip welded to the box it labels — otherwise a page break
    # can leave an unlabelled yellow rectangle at the top of the next page.
    hp.paragraph_format.keep_with_next = True
    for r in t.rows:
        trPr = r._tr.get_or_add_trPr()
        cant = OxmlElement('w:cantSplit')
        trPr.append(cant)
    run(hp, '✎  ', size=11, font='Segoe UI Symbol')
    run(hp, 'YOUR ANSWER', bold=True, size=10, color=RGBColor(0x7A, 0x5A, 0x00))

    body = t.rows[1].cells[0]
    body.text = ''
    cell_shade(body, 'FFFBEF')
    cell_borders(body, color='D9A400', sz=8)
    cell_margins(body, top=120, bottom=120)
    bp = body.paragraphs[0]
    bp.paragraph_format.space_after = Pt(0)
    run(bp, hint, italic=True, color=RGBColor(0xA0, 0x92, 0x6A))
    for _ in range(lines - 1):
        extra = body.add_paragraph()
        extra.paragraph_format.space_after = Pt(0)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    return t


def set_widths(table, widths):
    """Word ignores cell widths unless the table is fixed-layout with a grid."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    for tag in ('w:tblLayout', 'w:tblW'):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    insert_ordered(tblPr, layout)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), str(int(sum(widths) * 1440)))
    tw.set(qn('w:type'), 'dxa')
    insert_ordered(tblPr, tw)

    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w * 1440)))
        grid.append(gc)
    tblPr.addnext(grid)

    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def _data_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = ''
        cell_shade(c, 'E8EDF5')
        cell_margins(c, top=70, bottom=70)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        run(p, h, bold=True, size=10, color=NAVY)
    trPr = hdr._tr.get_or_add_trPr()
    hdr_repeat = OxmlElement('w:tblHeader')
    trPr.append(hdr_repeat)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            c.text = ''
            cell_margins(c, top=60, bottom=60)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if isinstance(val, tuple):
                text, bold = val
            else:
                text, bold = val, False
            run(p, text, bold=bold, size=10)
    for r in t.rows:
        r._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
    if widths:
        set_widths(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def _question(doc, num, title, ref, blocking):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    _rule(doc)
    h = doc.add_heading(level=2)
    run(h, '%s — %s' % (num, title), bold=True, size=14, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run(p, 'Spec reference: ', size=10, color=GREY)
    run(p, ref, size=10, bold=True, color=GREY)
    run(p, '     ', size=10)
    if blocking:
        run(p, 'BLOCKING', bold=True, size=10, color=RED)
    else:
        run(p, 'Not blocking', bold=True, size=10, color=GREY)


def _footer_page_numbers(doc, label):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run(p, f'{label}      ', size=9, color=GREY)
        r = p.add_run()
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        for instr, kind in (('begin', 'fldChar'), ('PAGE', 'instrText'), ('end', 'fldChar')):
            e = OxmlElement('w:' + kind)
            if kind == 'fldChar':
                e.set(qn('w:fldCharType'), instr)
            else:
                e.set(qn('xml:space'), 'preserve')
                e.text = ' PAGE '
            r._r.append(e)


def _base_styles(doc):
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.12
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    for name, size, color, before, after in (
        ('Heading 1', 20, NAVY, 0, 10),
        ('Heading 2', 14, NAVY, 20, 8),
        ('Heading 3', 11.5, NAVY, 12, 4),
    ):
        st = doc.styles[name]
        st.font.name = 'Calibri'
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.font.italic = False
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True



class Sheet:
    """A question sheet under construction. Every method appends to the document."""

    def __init__(self, title, subtitle, meta_rows, footer_label):
        self.doc = Document()
        _base_styles(self.doc)
        self._title(title, subtitle, meta_rows)
        self._footer_label = footer_label

    # ---- structure -------------------------------------------------------
    def _title(self, title, subtitle, meta_rows):
        d = self.doc
        h = d.add_heading(level=1)
        run(h, title, bold=True, size=22, color=NAVY)
        sub = d.add_paragraph()
        sub.paragraph_format.space_after = Pt(10)
        run(sub, subtitle, color=GREY, size=10.5)

        meta = d.add_table(rows=len(meta_rows), cols=2)
        no_table_borders(meta)
        set_widths(meta, [1.15, 5.3])
        for i, (k, v) in enumerate(meta_rows):
            kc, vc = meta.rows[i].cells
            for c in (kc, vc):
                c.text = ''
                cell_margins(c, top=30, bottom=30, left=0, right=80)
            kp = kc.paragraphs[0]
            kp.paragraph_format.space_after = Pt(0)
            run(kp, k, bold=True, size=10, color=NAVY)
            vp = vc.paragraphs[0]
            vp.paragraph_format.space_after = Pt(0)
            run(vp, v, size=10)
        d.add_paragraph().paragraph_format.space_after = Pt(6)

    def heading(self, text, level=2):
        self.doc.add_heading(text, level=level)

    def question(self, num, title, ref, blocking):
        _question(self.doc, num, title, ref, blocking)

    # ---- blocks ----------------------------------------------------------
    def para(self, parts, space_after=8, indent=0.0):
        return _para(self.doc, parts, space_after, indent)

    def quote(self, text):
        return _quote(self.doc, text)

    def bullet(self, parts, level=0):
        return _bullet(self.doc, parts, level)

    def numbered(self, parts):
        return _numbered(self.doc, parts)

    def option(self, label, rest='', level=0):
        return _option(self.doc, label, rest, level)

    def rule(self):
        return _rule(self.doc)

    def answer_box(self, hint='Click here and type your answer.', lines=4):
        return _answer_box(self.doc, hint, lines)

    def table(self, headers, rows, widths=None):
        return _data_table(self.doc, headers, rows, widths)

    def spacer(self, pt=6):
        self.doc.add_paragraph().paragraph_format.space_after = Pt(pt)

    def signoff(self):
        self.spacer()
        self.rule()
        t = self.doc.add_table(rows=2, cols=2)
        no_table_borders(t)
        set_widths(t, [1.3, 5.15])
        for i, (k, v) in enumerate((('Answered by', '\u2026' * 20), ('Date', '\u2026' * 20))):
            kc, vc = t.rows[i].cells
            for c in (kc, vc):
                c.text = ''
                cell_margins(c, top=60, bottom=60, left=0, right=80)
            kp = kc.paragraphs[0]
            kp.paragraph_format.space_after = Pt(0)
            run(kp, k, bold=True, size=10.5, color=NAVY)
            vp = vc.paragraphs[0]
            vp.paragraph_format.space_after = Pt(0)
            run(vp, v, size=10.5, color=GREY)

    def save(self, path):
        _footer_page_numbers(self.doc, self._footer_label)
        self.doc.save(path)
        return path
