# -*- coding: utf-8 -*-
"""Build the Order open-questions sheet as a .docx for a non-technical reviewer."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
RED = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x60, 0x60, 0x60)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()

# ---------- page + base styles ----------
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = BLACK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.12
normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

for name, size, color, before, after in (
    ('Heading 1', 20, NAVY, 0, 10),
    ('Heading 2', 14, NAVY, 20, 8),
    ('Heading 3', 11.5, NAVY, 12, 4),
):
    st = doc.styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.font.italic = False
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True


# ---------- OOXML schema ordering (LibreOffice/Word reject out-of-order children) ----------
_ORDERS = {
    'tcPr': ['w:cnfStyle', 'w:tcW', 'w:gridSpan', 'w:hMerge', 'w:vMerge', 'w:tcBorders',
             'w:shd', 'w:noWrap', 'w:tcMar', 'w:textDirection', 'w:tcFitText', 'w:vAlign',
             'w:hideMark'],
    'pPr': ['w:pStyle', 'w:keepNext', 'w:keepLines', 'w:pageBreakBefore', 'w:framePr',
            'w:widowControl', 'w:numPr', 'w:suppressLineNumbers', 'w:pBdr', 'w:shd', 'w:tabs',
            'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap', 'w:overflowPunct',
            'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd',
            'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents',
            'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
            'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr',
            'w:sectPr', 'w:pPrChange'],
    'tblPr': ['w:tblStyle', 'w:tblpPr', 'w:tblOverlap', 'w:bidiVisual', 'w:tblStyleRowBandSize',
              'w:tblStyleColBandSize', 'w:tblW', 'w:jc', 'w:tblCellSpacing', 'w:tblInd',
              'w:tblBorders', 'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook',
              'w:tblCaption', 'w:tblDescription', 'w:tblPrChange'],
    'rPr': ['w:rStyle', 'w:rFonts', 'w:b', 'w:bCs', 'w:i', 'w:iCs', 'w:caps', 'w:smallCaps',
            'w:strike', 'w:dstrike', 'w:outline', 'w:shadow', 'w:emboss', 'w:imprint',
            'w:noProof', 'w:snapToGrid', 'w:vanish', 'w:webHidden', 'w:color', 'w:spacing',
            'w:w', 'w:kern', 'w:position', 'w:sz', 'w:szCs', 'w:highlight', 'w:u', 'w:effect',
            'w:bdr', 'w:shd', 'w:fitText', 'w:vertAlign', 'w:rtl', 'w:cs', 'w:em', 'w:lang',
            'w:eastAsianLayout', 'w:specVanish', 'w:oMath'],
}


def _local(tag):
    return tag.split('}')[-1]


def insert_ordered(parent, child):
    """Insert child into parent at its schema-mandated position."""
    order = _ORDERS.get(_local(parent.tag))
    name = 'w:' + _local(child.tag)
    if order is None or name not in order:
        parent.append(child)
        return child
    idx = order.index(name)
    for existing in parent:
        en = 'w:' + _local(existing.tag)
        if en in order and order.index(en) > idx:
            existing.addprevious(child)
            return child
    parent.append(child)
    return child


# ---------- low-level helpers ----------
def shade(el, hexfill):
    for old in el.findall(qn('w:shd')):
        el.remove(old)
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hexfill)
    insert_ordered(el, sh)


def cell_shade(cell, hexfill):
    shade(cell._tc.get_or_add_tcPr(), hexfill)


def cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        e = OxmlElement('w:' + tag)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    insert_ordered(tcPr, mar)


def borders(el, color='BFBFBF', sz=6, edges=('top', 'left', 'bottom', 'right'), val='single'):
    """el is a tcPr or pPr-borders parent; returns the created element name used."""
    tag = 'w:tcBorders' if _local(el.tag) == 'tcPr' else 'w:pBdr'
    bd = el.find(qn(tag))
    if bd is None:
        bd = OxmlElement(tag)
        insert_ordered(el, bd)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge not in edges:
            continue
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), val)
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '4' if tag == 'w:pBdr' else '0')
        e.set(qn('w:color'), color)
        bd.append(e)


def cell_borders(cell, color='BFBFBF', sz=6, edges=('top', 'left', 'bottom', 'right')):
    borders(cell._tc.get_or_add_tcPr(), color, sz, edges)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    bd = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        bd.append(e)
    insert_ordered(tblPr, bd)


def run(p, text, bold=False, italic=False, color=None, size=None, font=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    if font is not None:
        r.font.name = font
        r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return r


def rich(p, parts):
    """parts: list of (text, bold, italic) or plain str."""
    for part in parts:
        if isinstance(part, str):
            run(p, part)
        else:
            text, bold, italic = (list(part) + [False, False])[:3]
            run(p, text, bold=bold, italic=italic)
    return p


def para(parts, space_after=8, indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    rich(p, parts if isinstance(parts, list) else [parts])
    return p


def quote(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.28)
    pf.right_indent = Inches(0.2)
    pf.space_before = Pt(6)
    pf.space_after = Pt(10)
    borders(p._p.get_or_add_pPr(), color='9DB2CE', sz=18, edges=('left',))
    shade(p._p.get_or_add_pPr(), 'F2F5FA')
    run(p, text, italic=True, color=GREY)
    return p


def bullet(parts, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + 0.28 * level)
    p.paragraph_format.space_after = Pt(4)
    rich(p, parts if isinstance(parts, list) else [parts])
    return p


def numbered(parts):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(6)
    rich(p, parts if isinstance(parts, list) else [parts])
    return p


def option(label, rest='', level=0):
    """A tickable option line: ☐  **label** rest"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.32 + 0.3 * level)
    pf.first_line_indent = Inches(-0.32)
    pf.space_after = Pt(5)
    run(p, '☐', size=13, font='Segoe UI Symbol')
    run(p, '   ')
    run(p, label, bold=True)
    if rest:
        run(p, rest)
    return p


def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    borders(p._p.get_or_add_pPr(), color='D0D0D0', sz=6, edges=('bottom',))
    return p


def answer_box(hint='Click here and type your answer.', lines=4):
    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_widths(t, [6.45])

    head = t.rows[0].cells[0]
    head.text = ''
    cell_shade(head, 'FFE9A8')
    cell_borders(head, color='D9A400', sz=8)
    cell_margins(head, top=60, bottom=60)
    hp = head.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    # Keep the header strip welded to the box it labels — otherwise a page break
    # can leave an unlabelled yellow rectangle at the top of the next page.
    hp.paragraph_format.keep_with_next = True
    for r in t.rows:
        trPr = r._tr.get_or_add_trPr()
        cant = OxmlElement('w:cantSplit')
        trPr.append(cant)
    run(hp, '✎  ', size=11, font='Segoe UI Symbol')
    run(hp, 'YOUR ANSWER', bold=True, size=10, color=RGBColor(0x7A, 0x5A, 0x00))

    body = t.rows[1].cells[0]
    body.text = ''
    cell_shade(body, 'FFFBEF')
    cell_borders(body, color='D9A400', sz=8)
    cell_margins(body, top=120, bottom=120)
    bp = body.paragraphs[0]
    bp.paragraph_format.space_after = Pt(0)
    run(bp, hint, italic=True, color=RGBColor(0xA0, 0x92, 0x6A))
    for _ in range(lines - 1):
        extra = body.add_paragraph()
        extra.paragraph_format.space_after = Pt(0)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    return t


def set_widths(table, widths):
    """Word ignores cell widths unless the table is fixed-layout with a grid."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    for tag in ('w:tblLayout', 'w:tblW'):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    insert_ordered(tblPr, layout)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), str(int(sum(widths) * 1440)))
    tw.set(qn('w:type'), 'dxa')
    insert_ordered(tblPr, tw)

    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w * 1440)))
        grid.append(gc)
    tblPr.addnext(grid)

    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def data_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = ''
        cell_shade(c, 'E8EDF5')
        cell_margins(c, top=70, bottom=70)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        run(p, h, bold=True, size=10, color=NAVY)
    trPr = hdr._tr.get_or_add_trPr()
    hdr_repeat = OxmlElement('w:tblHeader')
    trPr.append(hdr_repeat)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            c.text = ''
            cell_margins(c, top=60, bottom=60)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if isinstance(val, tuple):
                text, bold = val
            else:
                text, bold = val, False
            run(p, text, bold=bold, size=10)
    for r in t.rows:
        r._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
    if widths:
        set_widths(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def question(num, title, ref, blocking):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    rule()
    h = doc.add_heading(level=2)
    run(h, '%s — %s' % (num, title), bold=True, size=14, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run(p, 'Spec reference: ', size=10, color=GREY)
    run(p, ref, size=10, bold=True, color=GREY)
    run(p, '     ', size=10)
    if blocking:
        run(p, 'BLOCKING', bold=True, size=10, color=RED)
    else:
        run(p, 'Not blocking', bold=True, size=10, color=GREY)


def footer_page_numbers():
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run(p, 'Order module — open questions      ', size=9, color=GREY)
        r = p.add_run()
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        for instr, kind in (('begin', 'fldChar'), ('PAGE', 'instrText'), ('end', 'fldChar')):
            e = OxmlElement('w:' + kind)
            if kind == 'fldChar':
                e.set(qn('w:fldCharType'), instr)
            else:
                e.set(qn('xml:space'), 'preserve')
                e.text = ' PAGE '
            r._r.append(e)


# =====================================================================
#  CONTENT
# =====================================================================

title = doc.add_heading(level=1)
run(title, 'Order module — open questions', bold=True, size=22, color=NAVY)

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(14)
run(sub, 'Nine questions about the LINKZ Order Behaviour handover. '
         'Please answer in the yellow boxes and send the file back.', color=GREY, size=10.5)
sub.paragraph_format.space_after = Pt(10)

meta = doc.add_table(rows=3, cols=2)
no_table_borders(meta)
set_widths(meta, [1.15, 5.3])
for i, (k, v) in enumerate((
    ('To', 'The author of the LINKZ Order Behaviour handover'),
    ('From', 'The team building the Order module in the LINKZ v4 prototype'),
    ('Date sent', '……………………'),
)):
    kc, vc = meta.rows[i].cells
    for c in (kc, vc):
        c.text = ''
        cell_margins(c, top=30, bottom=30, left=0, right=80)
    kp = kc.paragraphs[0]
    kp.paragraph_format.space_after = Pt(0)
    run(kp, k, bold=True, size=10, color=NAVY)
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_after = Pt(0)
    run(vp, v, size=10)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ---------- How to answer ----------
doc.add_heading('How to answer this document', level=2)

para('You do not need the codebase, the prototype, or a Figma account. Everything each '
     'question refers to is quoted inside the question itself.')

numbered([('Type into the yellow boxes. ', True),
          'Every question is followed by a box headed ',
          ('YOUR ANSWER', True),
          '. Click inside it, delete the grey placeholder, and type. The box grows as you write.'])
numbered([('Tick a box by replacing ', True), ('☐', True),
          (' with an ', True), ('X', True),
          '. Click just to the right of the ☐, press Backspace once, and type a capital X. '
          'If none of the options fit, tick nothing and write what should happen instead.'])
numbered([('Please do not skip any question. ', True),
          'If one does not matter to you, write ',
          ('no preference', False, True),
          ' — that is a genuinely useful answer and we will choose a sensible default.'])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(4)
rich(p, [('Four of the nine block work that is otherwise ready to start', True),
         ' — Q1 to Q4, marked ',
         ('BLOCKING', True),
         '. If you are short of time, answer those four first and send them back; the rest '
         'can follow. The other five are not blocking, but each one is currently a guess on '
         'our side, so a one-line confirmation is worth having.'])

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------- Index ----------
doc.add_heading('The nine questions at a glance', level=2)
data_table(
    ['#', 'Question', 'Spec', 'Blocking'],
    [
        ['Q1', 'What amount should the even-out closing invoice carry?', '§5.5', ('Yes', True)],
        ['Q2', 'Are there four settlement statuses or five?', '§9', ('Yes', True)],
        ['Q3', 'Where do the Settlement and Payments rows come from?', '§9', ('Yes', True)],
        ['Q4', 'Should the order list really start empty?', '§4', ('Yes', True)],
        ['Q5', 'What does the Status filter do on the Payments tab?', '§9', 'No'],
        ['Q6', 'Should the date range actually filter the table?', '§9', 'No'],
        ['Q7', 'Can a purchase invoice only ever be settled at checkout?', '§6', 'No'],
        ['Q8', 'Do line items stay editable after money is collected?', '§5.2', 'No'],
        ['Q9', 'Order Report design references (a correction to confirm)', '§9', 'No'],
    ],
    widths=[0.5, 3.9, 0.65, 1.05],
)


# ---------- Q1 ----------
question('Q1', 'What amount should the even-out closing invoice carry?', '§5.5', True)
para([('This is your own open question in the handover, still unanswered.', True),
      ' It blocks the Send Order dialogs.'])
para([('The situation. ', True),
      'An order is worth IDR 5.000.000. An invoice for IDR 3.000.000 has been raised and '
      'paid. The user then edits the order down to IDR 3.000.000 — exactly what was '
      'already paid — and sends it. The spec calls this the ',
      ('even-out', False, True),
      ' send: it commits the adjustment, issues a closing invoice, and completes the order.'])
para([('The problem. ', True),
      'The paid invoice already covers the whole adjusted total, so there is nothing left to '
      'bill. The closing invoice currently comes out at ',
      ('IDR 0,00', True),
      '. Making it non-zero would mean going back and changing the earlier IDR 3.000.000 '
      'invoice, which contradicts §3.4 of your own spec:'])
quote('An invoice row records the state at the moment it was issued and never changes '
      'afterwards, except for its status.')
para([('Options as we see them:', True)])
option('A — Keep it at IDR 0,00.',
       ' The closing invoice acts as a marker that the order was evened out and closed. '
       'Invoice history stays untouched.')
option('B — Do not issue a closing invoice at all',
       ' on an even-out send. Just commit the adjustment and complete the order.')
option('C — Something else',
       ' — describe it below, including which existing invoice may be changed after the '
       'fact and how that squares with §3.4.')
answer_box()

# ---------- Q2 ----------
question('Q2', 'Are there four settlement statuses or five?', '§9', True)
para([('The conflict. ', True),
      'Your note on the Order Report page lists ', ('five', True), ' conditions:'])
data_table(
    ['Condition (as written in your note)', 'Status named in the note'],
    [
        ['Sudah paid tapi belum masuk excel', 'Pending Payment'],
        ['Sudah masuk excel tapi belum settled', 'Pending Payment (Not Yet Paid)'],
        ['Sudah masuk excel dan sudah settled', 'Paid Settled'],
        ['Sudah masuk excel, sudah settled, tapi turns out fraud', 'Chargeback'],
        ['Sudah masuk excel, belum settle, tapi turns out fraud', 'Canceled'],
    ],
    widths=[4.0, 2.4],
)
para(['But the design frame itself shows only ', ('four', True), ' status chips — ',
      ('Settled, Pending, Charge Back, Cancelled', True),
      ' — so the first two conditions in the table above both appear to the user as a '
      'single ', ('Pending', True), '.'])
para([('What we built: ', True),
      'the four from the frame, because that is what a user actually sees on screen.'])
option('A — Four is correct.',
       ' The two "pending" conditions look the same to the user; the distinction is internal only.')
option('B — Five is correct.',
       ' Please give us the exact wording for the fifth chip, and tell us how a reader is '
       'meant to tell the two pending states apart at a glance.')
answer_box('Click here and type your answer. If you tick B, please give the exact chip labels you want.')

# ---------- Q3 ----------
question('Q3', 'Where do the Settlement and Payments rows come from?', '§9', True)
para('You recommended feeding both tables from real activity in the prototype:')
quote('Settling a sales invoice writes a Settlement row; paying a purchase invoice writes a '
      'Payment row carrying the method chosen at checkout.')
para('We have built both tables from fixed sample data for now. Two things are unclear '
     'before we connect them to real activity.')

doc.add_heading('3a. Please confirm the two ways a row gets created', level=3)
option('Correct —', ' a Settlement row is created only when a SALES invoice is settled.')
option('Correct —', ' a Payments row is created only when a PURCHASE invoice is paid at checkout.')
para([('If either is wrong, ', False), 'say what else should create a row, in the box below.'],
     space_after=8)

doc.add_heading('3b. What moves a settlement from Pending to Settled?', level=3)
para('If a Settlement row appears the moment a sales invoice is marked paid, it presumably '
     'starts at Pending. This prototype has no backend, so nothing can later come along and '
     'mark it Settled by itself.')
option('A — Leave them Pending.',
       ' New rows stay at Pending forever; Settled only ever appears on the sample rows we seeded.')
option('B — Move to Settled on a timer.', ' Please say how long.')
option('C — Move to Settled by a manual action', ' in the prototype. Please say what the '
       'action is and where it lives.')
answer_box()

# ---------- Q4 ----------
question('Q4', 'Should the order list really start empty?', '§4', True)
para([('This is the one we would most like your decision on', True),
      ', because it changes several screens either way.'])
para([('What the spec says:', True)])
quote('Starts empty; the empty state is the front door with a single Create Order call to action.')
para([('Why that conflicts with the rest of the prototype. ', True),
      'This build is a ', ('training mock', True),
      '. Several screens that are already built read from pre-loaded orders:'])
bullet([('the Dashboard', True), ' shows order counts, GMV and revenue totals;'])
bullet([('the Finance screens', True), ' (Seller and Buyer Pay Later) reference invoices and orders;'])
bullet([('Checkout', True), ' is reached from a pre-loaded purchase order.'])
para('If the order list starts genuinely empty, those screens either show zeros, or show '
     'figures that no longer correspond to anything in the list.', space_after=10)
option('A — Start empty and accept the knock-on.',
       ' The list is the front door; the dashboard and finance screens show zeros until the '
       'trainee creates orders. Most faithful to §4.')
option('B — Keep the pre-loaded orders.',
       ' The empty state appears only after a trainee clears the demo data. Most useful for a '
       'training walkthrough.')
option('C — Pre-load the other screens separately',
       ' from the order list, so the list starts empty while the dashboard still demonstrates '
       'populated figures. This means the numbers will not reconcile against the list — '
       'is that acceptable in a mock?')
answer_box()

# ---------- Q5 ----------
question('Q5', 'What does the Status filter do on the Payments tab?', '§9', False)
para(['The Payments frame shows an ', ('All Status', True),
      ' filter in the toolbar. But the Payments table has no Status column — and your own '
      'comparison table says so explicitly:'])
quote('Status column — Settlement: Yes (Settled / Pending / Charge Back / Cancelled). '
      'Payments: No.')
para([('What we built: ', True),
      'the filter is hidden on the Payments tab and shown on Settlement only.'])
option('A — Correct, drop it', ' from the Payments tab.')
option('B — Keep it', ' — it filters on something that is not shown in the table. '
       'Please name the field below.')
answer_box()

# ---------- Q6 ----------
question('Q6', 'Should the date range actually filter the table?', '§9', False)
para('The date picker is built and applies a range correctly. Two loose ends remain.')
para([('6a. ', True), 'The sample rows are dated ', ('June 2026', True),
      ', but the default preset in the frame is ', ('Last 7 days', True),
      '. If the range filtered literally against today’s date, the table would open ',
      ('empty', True), ' — which is not what the frame shows.'])
para([('6b. ', True),
      'Right now the chosen range is applied and displayed, but it does not actually filter '
      'the sample rows.'], space_after=10)
option('A — Do not filter in the mock.',
       ' The picker demonstrates the interaction; the rows stay put. This is what we have now.')
option('B — Filter for real',
       ' — and either move the sample data inside the default range, or change the default '
       'preset. Please say which.')
answer_box()

# ---------- Q7 ----------
question('Q7', 'Can a purchase invoice only ever be settled at checkout?', '§6', False)
para('Your table gives the two invoice types different row actions:')
data_table(
    ['', 'Sales invoice', 'Purchase invoice'],
    [
        ['Row actions',
         'Menu: Download PDF · Mark as Paid · Void Invoice',
         ('Download icon only — no menu', True)],
    ],
    widths=[1.1, 2.7, 2.6],
)
para([('The consequence. ', True),
      'A purchase invoice can then only be closed by ', ('paying it at checkout', True),
      ' — there is no Mark as Paid and no Void. So a purchase invoice raised in error '
      'cannot be voided at all, and the order it belongs to cannot be closed without paying it.'],
     space_after=10)
option('A — That is intended.', ' Purchase invoices are settled only at checkout; that is the point.')
option('B — Not intended.', ' Purchase invoices should also offer (tick what applies):')
option('Void Invoice', level=1)
option('Mark as Paid', level=1)
option('Something else', ' — describe it below.', level=1)
answer_box()

# ---------- Q8 ----------
question('Q8', 'Do line items stay editable after money has been collected?', '§5.2', False)
quote('Rows stay editable after the order is sent, until it is Completed or Cancelled. '
      'Only Buyer/Seller Info locks on send.')
para(['Read literally, that includes an order with a ', ('fully paid', True),
      ' invoice against it. We believe that is deliberate — it is exactly what makes the '
      'Overpaid state in §5.4 reachable, since overpaid arises only by ',
      ('reducing an order after money has been collected', False, True), '.'])
para('We have implemented it that way. We are asking you to confirm, because it is the kind '
     'of rule that looks like a bug to a reviewer who has not read §5.4.', space_after=10)
option('A — Confirmed.',
       ' Editing stays open until the order is Completed or Cancelled, even when a paid invoice '
       'sits against it.')
option('B — No', ' — the product table should lock at some earlier point. Please say when.')
answer_box()

# ---------- Q9 ----------
question('Q9', 'Order Report design references — a correction to confirm', '§9', False)
para(['The three design references cited in §9 — ',
      ('4072:99011, 4072:99211, 4072:99053', True),
      ' — ', ('do not exist', True),
      ' in the Figma file we were given. We found the Order Report elsewhere in that file and '
      'built against these frames instead:'])
data_table(
    ['Screen', 'Frame', 'Link'],
    [
        ['Settlement', '7017:1308',
         'figma.com/design/eX8Lc53tVFuY2QEDW4t1QT/…?node-id=7017-1308'],
        ['Payments', '7017:1508',
         'figma.com/design/eX8Lc53tVFuY2QEDW4t1QT/…?node-id=7017-1508'],
        ['Date picker, open', '7017:1350',
         'figma.com/design/eX8Lc53tVFuY2QEDW4t1QT/…?node-id=7017-1350'],
    ],
    widths=[1.35, 1.05, 4.0],
)
option('A — Confirmed.',
       ' The 4072 references came from an earlier file; the frames above are the intended design.')
option('B — No',
       ' — the 4072 frames are newer and live somewhere else. Please send us the file link.')
answer_box()

# ---------- Anything else ----------
doc.add_paragraph().paragraph_format.space_after = Pt(0)
rule()
h = doc.add_heading(level=2)
run(h, 'Anything else', bold=True, size=14, color=NAVY)
para('If reviewing these surfaced something we have not asked about — a rule you expected '
     'to see questioned, or a decision you would like revisited — please put it here.')
answer_box('Optional.', lines=5)

# ---------- Sign-off ----------
doc.add_paragraph().paragraph_format.space_after = Pt(6)
rule()
sign = doc.add_table(rows=2, cols=2)
no_table_borders(sign)
set_widths(sign, [1.3, 5.15])
for i, (k, v) in enumerate((('Answered by', '………………………………………………'),
                            ('Date', '………………………………………………'))):
    kc, vc = sign.rows[i].cells
    for c in (kc, vc):
        c.text = ''
        cell_margins(c, top=60, bottom=60, left=0, right=80)
    kp = kc.paragraphs[0]
    kp.paragraph_format.space_after = Pt(0)
    run(kp, k, bold=True, size=10.5, color=NAVY)
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_after = Pt(0)
    run(vp, v, size=10.5, color=GREY)

footer_page_numbers()

out = '/home/user/Linkz-v4-landing-mock/docs/Order module - open questions.docx'
doc.save(out)
print('saved:', out)
